import os
import re
import PyPDF2 as pdf
import docx
from dotenv import load_dotenv
import fitz
import google.generativeai as genai
import pdfplumber
import csv
from PyPDF2 import PdfReader
from docx import Document
from docx2pdf import convert
import env
import shutil
import csv

# Set up Google API key and configure Generative AI
genai.configure(api_key=env.GOOGLE_API_KEY) 
#----------------------------------------------------------------

#GEMINI FOR CANDIDATE
class Gemini:

    RESUME_DIRECTORY = "C:\\xampp\\htdocs\\login_signup\\uploads\\Candidate\\CV\\resume.pdf"
    JOB_DESC_DIRECTORY = "C:\\xampp\\htdocs\\login_signup\\uploads\\Candidate\\JD\\jd.txt"
    
    @staticmethod
    def get_gemini_response(prompt):
        model = genai.GenerativeModel(env.MODEL)
        response = model.generate_content(prompt)
        return response.text

    @staticmethod
    def extract_text_from_pdf(resume_file_path):
        """Extract text from the uploaded PDF resume."""
        with open(resume_file_path, 'rb') as file:
            reader = pdf.PdfReader(file)
            text = ""
            for page in range(len(reader.pages)):
                page_text = reader.pages[page].extract_text()
                text += str(page_text)
        return text

    @staticmethod
    def load_job_description(job_description_file_path):
        """Load job description from the text file."""
        with open(job_description_file_path, 'r') as file:
            return file.read()

    @staticmethod
    def process_resume():
        job_desc = Gemini.load_job_description(Gemini.JOB_DESC_DIRECTORY)
        resume_text = Gemini.extract_text_from_pdf(Gemini.RESUME_DIRECTORY)
        
        """Process the resume based on the user's selected action."""
        
        prompt = f"""
        You are an expert in the field of HR. You have been asked to evaluate a candidate for a job position. The job description and the resume of the candidate are provided below. You will first show the percentage match between the job description and CV in one line and write the main key skills and project missing from the CV for that job write it in short points. Then create viva questions that could be asked based on the skills, job role, previous experience, and projects mentioned in the CV. Also, generate technical questions from CV projects and experience.

        Job Description:
        {job_desc}

        Resume:
        {resume_text}
        """
        
        response_text = Gemini.get_gemini_response(prompt)
        return response_text

class GeminiHR:
    RESUME_HR_DIRECTORY = "C:\\xampp\\htdocs\\login_signup\\uploads\\HR\\CV\\"
    JOB_DESC_HR_DIRECTORY = "C:\\xampp\\htdocs\\login_signup\\uploads\\HR\\JD\\jd.txt"

    @staticmethod
    def get_gemini_response(prompt):
        """Get response from Gemini AI."""
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(prompt)
        return response.text

    @staticmethod
    def extract_pdf_text(file_path):
        """Extract text from PDF file."""
        with pdfplumber.open(file_path) as pdf:
            return "".join(page.extract_text() for page in pdf.pages)

    @staticmethod
    def extract_docx_text(file_path):
        """Extract text from DOCX file."""
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    @staticmethod
    def process_resume():
        """Process the resumes based on the user's selected action."""
        RESUME_DIRECTORY = "C:\\xampp\\htdocs\\login_signup\\uploads\\HR\\CV\\"
        JOB_DESC_DIRECTORY = "C:\\xampp\\htdocs\\login_signup\\uploads\\HR\\JD\\jd.txt"

        job_desc = Gemini.load_job_description(JOB_DESC_DIRECTORY)

        pdf_files = [
            file for file in os.listdir(RESUME_DIRECTORY)
            if file.lower().endswith('.pdf')
        ]

        percentage_mapping = {}

        for pdf_file in pdf_files:
            pdf_path = os.path.join(RESUME_DIRECTORY, pdf_file)
            
            resume_text = Gemini.extract_text_from_pdf(pdf_path)

            prompt = f"""
            You are an expert in the field of HR. You have been asked to evaluate a candidate for a job position. The job description and the resume of the candidate are provided below. Your task is to show the percentage match between the job description and CV in one line. Just write the number without using %.
            Job Description:
            {job_desc}

            Resume:
            {resume_text}
            """
            
            response = Gemini.get_gemini_response(prompt).strip()

            try:
                percentage = int(response.replace("\n", "").strip())
            except ValueError:
                percentage = 0 
            
            percentage_mapping[pdf_file] = percentage
        percentage_mapping = dict(sorted(percentage_mapping.items(), key=lambda item: item[1], reverse=True))
        return percentage_mapping

#----------------------------------------------------------------
# Extract information from resumes and save to CSV
#----------------------------------------------------------------

HR_RESUME_DIRECTORY = "C:\\xampp\\htdocs\\login_signup\\uploads\\HR\\CV\\"

class ResumeExtractor:

    @staticmethod
    def extract_emails(text):
        """Extract email addresses from text."""
        return re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)

    @staticmethod
    def extract_phone_numbers(text):
        """Extract phone numbers from text."""
        return re.findall(r'\+?880[-\s]?\d{4}[-\s]?\d{6}|\d{5}[-\s]?\d{6}', text)

    @staticmethod
    def extract_text_from_pdf(resume_file_path):
        """Extract text from the uploaded PDF resume using pdfplumber."""
        text = ""
        try:
            with pdfplumber.open(resume_file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:  # Only add text if extraction is successful
                        text += page_text
        except Exception as e:
            print(f"Error extracting text from PDF {resume_file_path}: {e}")
        return text

    # Static method to extract emails
    @staticmethod
    def extract_emails(text):
        return re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zAZ0-9.-]+', text)

    # Static method to extract phone numbers
    @staticmethod
    def extract_phone_numbers(text):
        return re.findall(r'\+?880[-\s]?\d{4}[-\s]?\d{6}|\d{5}[-\s]?\d{6}', text)

    # Static method to read content from a file (modified for PDF handling)
    @staticmethod
    def read_file_content(file_path):
        """Read file content, handling both PDFs and text files."""
        try:
            # Check if the file is a PDF based on the extension
            if file_path.lower().endswith('.pdf'):
                return ResumeExtractor.extract_text_from_pdf(file_path)  # Extract text from PDF
            else:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()  # Read text files as usual
        except Exception as e:
            print(f"Error reading file {file_path}: {e}")
            return ""

    # Static method to extract resume information
    @staticmethod
    def extract_resume_info(fileNames):

        resumes = []
        OUTPUT_DIRECTORY = "./output"

        # Read all files in the HR resume directory
        for file_name in fileNames:
            file_path = os.path.join(HR_RESUME_DIRECTORY, file_name)
            if os.path.isfile(file_path):  # Ensure it's a file
                content = ResumeExtractor.read_file_content(file_path)
                if content:
                    resumes.append((file_name, content))

        info = {}

        for name, text in resumes:
            emails = ResumeExtractor.extract_emails(text)
            phones = ResumeExtractor.extract_phone_numbers(text)
            
            # Store additional emails and phones as references
            info[name] = {
                "candidate_email" : emails[0] if emails else "No email found",
                "candidate_phone" : phones[0] if phones else "No phone number found",
                "candidate_Ref_emails": emails[1:] if len(emails) > 1 else [],
                "candidate_Ref_phones": phones[1:] if len(phones) > 1 else []
            }
            csv_file_path = os.path.join(OUTPUT_DIRECTORY, "resume_info.csv")

            with open(csv_file_path, mode='w', newline='', encoding='utf-8') as csv_file:
                fieldnames = ["Name", "Email", "Phone", "Reference Emails", "Reference Phones"]
                writer = csv.DictWriter(csv_file, fieldnames=fieldnames)

                writer.writeheader()
                for file_name, details in info.items():
                    writer.writerow({
                        "Name": file_name,
                        "Email": details["candidate_email"],
                        "Phone": details["candidate_phone"],
                        "Reference Emails": ", ".join(details["candidate_Ref_emails"]),
                        "Reference Phones": ", ".join(details["candidate_Ref_phones"])
                    })
        return info

#----------------------------------------------------------------
# Generate questions based on resumes and job description
#----------------------------------------------------------------

class GeminiHR_Generate_questions:
    RESUME_DIR = "C:\\xampp\\htdocs\\login_signup\\uploads\\HR\\CV\\"
    JD_PATH = "C:\\xampp\\htdocs\\login_signup\\uploads\\Candidate\\JD\\jd.txt"
    OUTPUT_DIR = "C:\\xampp\\htdocs\\login_signup\\output\\"

    @staticmethod
    def get_gemini_response(prompt):
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(prompt)
        return response.text

    @staticmethod
    def extract_text_from_pdf(path):
        with open(path, 'rb') as file:
            reader = pdf.PdfReader(file)
            return ''.join([page.extract_text() for page in reader.pages if page.extract_text()])

    @staticmethod
    def load_job_description(path):
        with open(path, 'r') as file:
            return file.read()

    @staticmethod
    def save_text_as_pdf(text, output_path):
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), text, fontsize=11)
        doc.save(output_path)
        doc.close()

    @staticmethod
    def save_text_as_docx(text, output_path):
        doc = Document()
        for line in text.split('\n'):
            doc.add_paragraph(line)
        doc.save(output_path)

    @staticmethod
    def process_resumes():
        if not os.path.exists(GeminiHR_Generate_questions.OUTPUT_DIR):
            os.makedirs(GeminiHR_Generate_questions.OUTPUT_DIR)

        jd_text = GeminiHR_Generate_questions.load_job_description(GeminiHR_Generate_questions.JD_PATH)

        for filename in os.listdir(GeminiHR_Generate_questions.RESUME_DIR):
            if filename.lower().endswith('.pdf'):
                resume_path = os.path.join(GeminiHR_Generate_questions.RESUME_DIR, filename)
                resume_text = GeminiHR_Generate_questions.extract_text_from_pdf(resume_path)

                prompt = f"""
                You are a Highly Expert HR. Your main task is to go through the job description throughly. Read that nicely and understand that nicely. Then go through the resume of the candidate. Read that nicely and understand that nicely. Then you have to generate the questions that could be asked based on the skills, job role, previous experience, and projects mentioned in the CV. Also, generate technical questions from CV projects and experience. The Question must be of high Standard and should be related to the job role and the skills mentioned in the CV. Also ask some Advance level questions based on Skills and project which relate to the job description. No need to ask general questions. The total number question will be between 10-15

                Job Description:
                {jd_text}

                Resume:
                {resume_text}
                """

                questions = GeminiHR_Generate_questions.get_gemini_response(prompt)
                # Create filenames
                base_name = os.path.splitext(filename)[0]
                docx_output_path = os.path.join(GeminiHR_Generate_questions.OUTPUT_DIR, base_name + "_questions.docx")
                # Save both versions
                GeminiHR_Generate_questions.save_text_as_docx(questions, docx_output_path)
                print(f"Generated: {base_name}questions.docx")


